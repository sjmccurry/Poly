import os
import re
import json
import tkinter as tk
from tkinter import filedialog, messagebox, Entry, IntVar, Checkbutton
from docx import Document

CONFIG_FILE = "poly_config.json"
LOCAL_VERSION_FILE = "version.txt"
DEFAULT_TEMPLATE_FOLDER = "templates"

COLORS = {
    "bg_dark": "#121212",
    "bg_panel": "#1e1e1e",
    "bg_header": "#1f1f1f",
    "bg_input": "#2a2a2a",
    "bg_button": "#3a3a3a",
    "bg_button_hover": "#505050",
    "accent": "#0078d4",
    "accent_hover": "#005ea2",
    "fg": "white",
}

def make_button(parent, text, command, primary=False):
    bg = COLORS["accent"] if primary else COLORS["bg_button"]
    bg_hover = COLORS["accent_hover"] if primary else COLORS["bg_button_hover"]
    fg = "#ffffff" if primary else "#d0d0d0"

    lbl = tk.Label(
        parent, text=text, bg=bg, fg=fg,
        font=("Segoe UI", 10, "bold"),
        padx=10, pady=8, cursor="hand2",
        anchor="w", width=20,
    )
    lbl.bind("<Button-1>", lambda e: command())
    lbl.bind("<Enter>", lambda e: lbl.config(bg=bg_hover, fg="#ffffff"))
    lbl.bind("<Leave>", lambda e: lbl.config(bg=bg, fg=fg))
    return lbl


def read_version():
    if os.path.exists(LOCAL_VERSION_FILE):
        with open(LOCAL_VERSION_FILE) as f:
            return f.read().strip()
    return "vUnknown"


def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE) as f:
                data = json.load(f)
            return (
                data.get("template_folder", DEFAULT_TEMPLATE_FOLDER),
                data.get("output_folder", os.getcwd()),
            )
        except Exception:
            pass
    return DEFAULT_TEMPLATE_FOLDER, os.getcwd()


def save_config(template_folder, output_folder):
    try:
        with open(CONFIG_FILE, "w") as f:
            json.dump({"template_folder": template_folder, "output_folder": output_folder}, f)
    except Exception:
        pass


def extract_placeholders(doc):
    pattern = r"\{poly\.([a-zA-Z0-9_ ]+)\}"
    found = set()
    for para in doc.paragraphs:
        found.update(re.findall(pattern, para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                found.update(re.findall(pattern, cell.text))
    return found


def field_label(field_name):
    pretty = re.sub(r"([a-z])([A-Z])", r"\1 \2", field_name)
    return pretty.replace("_", " ")


def get_run_font(run):
    if run.font.name:
        return run.font.name
    try:
        rfonts = run._element.xpath(".//w:rFonts")
        if rfonts:
            ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            return rfonts[0].get(f"{ns}ascii")
    except Exception:
        pass
    return None


def char_attrs(run):
    return {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "font": get_run_font(run),
        "size": run.font.size,
    }


def runs_to_chars(paragraph):
    chars = []
    for run in paragraph.runs:
        attrs = char_attrs(run)
        for ch in run.text:
            chars.append({"char": ch, **attrs})
    return chars


def apply_replacements_to_paragraph(paragraph, replacements):
    pattern = re.compile(r"\{poly\.([a-zA-Z0-9_ ]+)\}")
    chars = runs_to_chars(paragraph)
    full_text = "".join(c["char"] for c in chars)

    matches = list(pattern.finditer(full_text))
    if not matches:
        return

    new_chars = []
    cursor = 0
    for match in matches:
        start, end = match.span()
        field = match.group(1)
        replacement = replacements.get(field, match.group(0))

        new_chars.extend(chars[cursor:start])

        fmt = chars[start] if start < len(chars) else {
            "bold": None, "italic": None, "underline": None, "font": None, "size": None
        }
        for ch in replacement:
            new_chars.append({"char": ch, **{k: fmt[k] for k in ("bold", "italic", "underline", "font", "size")}})

        cursor = end

    new_chars.extend(chars[cursor:])

    for run in paragraph.runs:
        run.text = ""

    active_run = None
    for ch in new_chars:
        needs_new_run = (
            active_run is None
            or active_run.bold != ch["bold"]
            or active_run.italic != ch["italic"]
            or active_run.underline != ch["underline"]
            or (ch["font"] and active_run.font.name != ch["font"])
            or (ch["size"] and active_run.font.size != ch["size"])
        )
        if needs_new_run:
            active_run = paragraph.add_run()
            active_run.bold = ch["bold"]
            active_run.italic = ch["italic"]
            active_run.underline = ch["underline"]
            if ch["font"]:
                try:
                    active_run.font.name = ch["font"]
                except Exception:
                    pass
            if ch["size"]:
                try:
                    active_run.font.size = ch["size"]
                except Exception:
                    pass
        active_run.text += ch["char"]


def fill_document(doc, replacements):
    for para in doc.paragraphs:
        apply_replacements_to_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    apply_replacements_to_paragraph(para, replacements)


def make_scrollable_frame(parent):
    canvas = tk.Canvas(parent, bg=COLORS["bg_panel"], highlightthickness=0)
    scrollbar = tk.Scrollbar(parent, orient="vertical", command=canvas.yview)
    frame = tk.Frame(canvas, bg=COLORS["bg_panel"])

    frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)

    scrollbar.pack(side="right", fill="y")
    canvas.pack(side="left", fill="both", expand=True)

    return frame, canvas


def bind_scroll_to_canvas(widget, canvas):
    def on_scroll(event):
        delta = -1 * int(event.delta / 120) if os.name == "nt" else -1 * int(event.delta)
        canvas.yview_scroll(delta, "units")

    def on_enter(_):
        widget.bind_all("<MouseWheel>", on_scroll)
        widget.bind_all("<Button-4>", lambda e: canvas.yview_scroll(-1, "units"))
        widget.bind_all("<Button-5>", lambda e: canvas.yview_scroll(1, "units"))

    def on_leave(_):
        widget.unbind_all("<MouseWheel>")
        widget.unbind_all("<Button-4>")
        widget.unbind_all("<Button-5>")

    widget.bind("<Enter>", on_enter)
    widget.bind("<Leave>", on_leave)


class PolyApp:
    def __init__(self, root):
        self.root = root
        self.root.title(f"Poly Document Generator {read_version()}")
        self.root.geometry("1280x720")
        self.root.configure(bg=COLORS["bg_dark"])

        self.selected_templates = []
        self.entries = {}

        self.template_folder, self.output_folder = load_config()

        self._build_ui()
        self.load_templates()

    def _build_ui(self):
        header = tk.Frame(self.root, bg=COLORS["bg_header"])
        header.pack(fill="x")
        tk.Label(
            header, text="Poly Document Generator",
            font=("Segoe UI", 24, "bold"),
            bg=COLORS["bg_header"], fg=COLORS["fg"]
        ).pack(pady=10)

        main = tk.Frame(self.root, bg=COLORS["bg_dark"])
        main.pack(fill="both", expand=True, padx=15, pady=10)

        self._build_template_panel(main)
        self._build_fields_panel(main)
        self._build_actions_panel(main)

    def _build_template_panel(self, parent):
        frame = tk.Frame(parent, bg=COLORS["bg_panel"], width=400)
        frame.pack(side="left", fill="y", padx=(0, 10))
        frame.pack_propagate(False)

        tk.Label(frame, text="Templates", bg=COLORS["bg_panel"], fg=COLORS["fg"],
                 font=("Segoe UI", 12, "bold")).pack(pady=10)

        self.template_frame, template_canvas = make_scrollable_frame(frame)
        bind_scroll_to_canvas(self.template_frame, template_canvas)

    def _build_fields_panel(self, parent):
        frame = tk.Frame(parent, bg=COLORS["bg_panel"])
        frame.pack(side="left", fill="both", expand=True)

        tk.Label(frame, text="Fill in Fields", bg=COLORS["bg_panel"], fg=COLORS["fg"],
                 font=("Segoe UI", 12, "bold")).pack(pady=10)

        self.fields_frame, fields_canvas = make_scrollable_frame(frame)
        bind_scroll_to_canvas(self.fields_frame, fields_canvas)

    def _build_actions_panel(self, parent):
        frame = tk.Frame(parent, bg=COLORS["bg_panel"], width=200)
        frame.pack(side="left", fill="y", padx=(10, 0))
        frame.pack_propagate(False)

        tk.Label(frame, text="Actions", bg=COLORS["bg_panel"], fg=COLORS["fg"],
                 font=("Segoe UI", 12, "bold")).pack(pady=10)

        make_button(frame, "Generate Document(s)", self.generate_documents, primary=True).pack(pady=10, padx=10, fill="x")
        make_button(frame, "Upload Template", self.upload_template).pack(pady=5, padx=10, fill="x")
        make_button(frame, "Select Template Folder", self.select_template_folder).pack(pady=5, padx=10, fill="x")

    def select_template_folder(self):
        folder = filedialog.askdirectory(title="Select Template Folder")
        if folder:
            self.template_folder = folder
            save_config(self.template_folder, self.output_folder)
            self.load_templates()

    def upload_template(self):
        src_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if not src_path:
            return
        dest_path = os.path.join(self.template_folder, os.path.basename(src_path))
        with open(src_path, "rb") as src, open(dest_path, "wb") as dst:
            dst.write(src.read())
        self.load_templates()

    def load_templates(self):
        for widget in self.template_frame.winfo_children():
            widget.destroy()

        os.makedirs(self.template_folder, exist_ok=True)

        for filename in sorted(f for f in os.listdir(self.template_folder) if f.endswith(".docx")):
            var = IntVar()
            cb = Checkbutton(
                self.template_frame, text=filename, variable=var,
                font=("Segoe UI", 10), anchor="w",
                command=self.update_selected_templates,
                bg=COLORS["bg_panel"], fg=COLORS["fg"],
                wraplength=280, justify="left",
                selectcolor=COLORS["bg_panel"],
                activebackground=COLORS["bg_panel"],
            )
            cb.var = var
            cb.filename = filename
            cb.pack(fill="x", anchor="w", padx=10, pady=2)

    def update_selected_templates(self):
        self.selected_templates = [
            w.filename for w in self.template_frame.winfo_children()
            if isinstance(w, Checkbutton) and w.var.get()
        ]
        self.load_fields()

    def load_fields(self):
        saved_values = {key: entry.get() for key, entry in self.entries.items()}

        for widget in self.fields_frame.winfo_children():
            widget.destroy()

        all_fields = set()
        for filename in self.selected_templates:
            try:
                doc = Document(os.path.join(self.template_folder, filename))
                all_fields.update(extract_placeholders(doc))
            except Exception:
                continue

        self.entries = {}
        for field in sorted(all_fields):
            tk.Label(
                self.fields_frame, text=field_label(field),
                bg=COLORS["bg_panel"], fg=COLORS["fg"],
                font=("Segoe UI", 10)
            ).pack(anchor="w", padx=10, pady=(10, 2))

            entry = Entry(
                self.fields_frame, width=50,
                bg=COLORS["bg_input"], fg=COLORS["fg"],
                insertbackground=COLORS["fg"],
                font=("Segoe UI", 10)
            )
            entry.pack(anchor="w", padx=10, pady=2)

            if field in saved_values:
                entry.insert(0, saved_values[field])

            self.entries[field] = entry

    def generate_documents(self):
        output_dir = filedialog.askdirectory(title="Select Output Folder", initialdir=self.output_folder)
        if not output_dir:
            return

        self.output_folder = output_dir
        save_config(self.template_folder, self.output_folder)

        values = {key: entry.get() for key, entry in self.entries.items()}

        for filename in self.selected_templates:
            try:
                doc = Document(os.path.join(self.template_folder, filename))
                fill_document(doc, values)
                out_name = filename.replace(".docx", "_filled.docx")
                doc.save(os.path.join(output_dir, out_name))
            except Exception as e:
                messagebox.showerror("Error", f"Could not process {filename}:\n{e}")
                return

        os.startfile(output_dir)
        messagebox.showinfo("Success", f"Documents saved to '{output_dir}'.")


if __name__ == "__main__":
    root = tk.Tk()
    app = PolyApp(root)
    root.mainloop()
